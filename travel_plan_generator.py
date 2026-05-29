# 没任何卵用的脚本

"""
旅游方案一键生成工具
整合功能：天气查询 + 交通价格查询 + 方案导出为Word文档
"""

import requests
import json
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn



# 天气查询（调用免费天气API）


class WeatherService:
    """天气查询服务，使用和风天气免费API"""

    # 和风天气免费API（需要注册获取key，或使用备用方案）
    BASE_URL = "https://devapi.qweather.com/v7"

    @staticmethod
    def get_weather(city_name: str, date_str: str = None) -> dict:
        """
        查询城市天气
        :param city_name: 城市名称，如"北京"
        :param date_str: 日期，如"2025-06-01"，不填则查询今天
        :return: 天气信息字典
        """
        # 方案1：使用和风天气API（需要API key）
        # 方案2：使用wttr.in（免费，无需key）
        try:
            url = f"https://wttr.in/{city_name}?format=j1&lang=zh"
            response = requests.get(url, timeout=10)
            data = response.json()

            current = data.get("current_condition", [{}])[0]
            weather = data.get("weather", [{}])[0]

            result = {
                "city": city_name,
                "temp_max": int(weather.get("maxtempC", 25)),
                "temp_min": int(weather.get("mintempC", 15)),
                "current_temp": int(current.get("temp_C", 20)),
                "desc": current.get("lang_zh", [{}])[0].get("value", "多云"),
                "humidity": current.get("humidity", "60"),
                "wind": current.get("windspeedKmph", "10"),
            }
            return result
        except Exception as e:
            print(f"天气查询失败: {e}")
            return {
                "city": city_name,
                "temp_max": 30,
                "temp_min": 18,
                "current_temp": 24,
                "desc": "查询失败，请手动确认",
                "humidity": "60",
                "wind": "10",
            }

    @staticmethod
    def get_clothing_advice(weather: dict) -> str:
        """根据天气生成穿衣建议"""
        temp_max = weather["temp_max"]
        temp_min = weather["temp_min"]

        advice = []

        if temp_max >= 35:
            advice.append("高温天气，注意防暑降温")
            advice.append("上衣：短袖、背心等清凉透气衣物")
            advice.append("下装：短裤、薄裙")
            advice.append("必备：防晒霜、遮阳帽、太阳镜、便携水杯")
        elif temp_max >= 28:
            advice.append("天气炎热")
            advice.append("上衣：短袖T恤、薄衬衫")
            advice.append("下装：短裤、薄款长裤")
            advice.append("必备：防晒霜、太阳镜、雨伞")
        elif temp_max >= 22:
            advice.append("天气舒适")
            advice.append("上衣：短袖+薄外套（早晚温差大）")
            advice.append("下装：长裤、牛仔裤")
            advice.append("必备：薄外套、雨伞")
        elif temp_max >= 15:
            advice.append("天气偏凉")
            advice.append("上衣：长袖T恤+卫衣/薄毛衣")
            advice.append("下装：长裤")
            advice.append("必备：厚外套、雨伞")
        else:
            advice.append("天气较冷")
            advice.append("上衣：保暖内衣+毛衣+厚外套/羽绒服")
            advice.append("下装：厚裤、加绒裤")
            advice.append("必备：围巾、手套、帽子")

        if temp_max - temp_min >= 10:
            advice.append(f"注意：昼夜温差达{temp_max - temp_min}°C，建议洋葱式穿衣")

        return "\n".join(advice)



# 交通价格查询（调用公开API）


class TransportService:
    """交通价格查询服务"""

    @staticmethod
    def search_train(from_city: str, to_city: str, date_str: str = None) -> list:
        """
        查询火车/高铁信息
        注意：12306官方API不对外开放，此处使用模拟数据或第三方接口
        建议实际使用时接入携程/去哪儿API
        """
        # 常见城市间高铁参考数据（实际项目应接入API）
        train_data = {
            ("天津", "北京"): [
                {"type": "高铁", "train_no": "C2025", "depart": "06:00", "arrive": "06:30",
                 "duration": "30分钟", "price": 54.5},
                {"type": "高铁", "train_no": "C2027", "depart": "07:00", "arrive": "07:34",
                 "duration": "34分钟", "price": 54.5},
                {"type": "高铁", "train_no": "G350", "depart": "08:00", "arrive": "08:30",
                 "duration": "30分钟", "price": 65.5},
            ],
            ("北京", "大连"): [
                {"type": "高铁", "train_no": "G3502", "depart": "07:00", "arrive": "12:00",
                 "duration": "5小时", "price": 400},
                {"type": "高铁", "train_no": "G390", "depart": "09:00", "arrive": "14:00",
                 "duration": "5小时", "price": 400},
            ],
            ("成都", "西安"): [
                {"type": "高铁", "train_no": "G350", "depart": "07:00", "arrive": "10:12",
                 "duration": "3小时12分", "price": 263},
                {"type": "高铁", "train_no": "G2204", "depart": "08:08", "arrive": "12:06",
                 "duration": "3小时58分", "price": 263},
            ],
        }

        key = (from_city, to_city)
        results = train_data.get(key, [])

        if not results:
            # 尝试反向查询
            key = (to_city, from_city)
            results = train_data.get(key, [])

        if not results:
            print(f"暂无 {from_city}→{to_city} 的火车数据，请手动查询12306")
            results = [{"type": "高铁", "train_no": "待查询", "depart": "--",
                        "arrive": "--", "duration": "待确认", "price": 0}]

        return results

    @staticmethod
    def search_flight(from_city: str, to_city: str, date_str: str = None) -> list:
        """
        查询机票信息
        注意：实际项目应接入携程/去哪儿/飞猪API
        """
        # 参考数据
        flight_data = {
            ("成都", "北京"): [
                {"airline": "国航", "flight_no": "CA4196", "depart": "07:00", "arrive": "09:35",
                 "duration": "2小时35分", "price": 500},
                {"airline": "川航", "flight_no": "3U8883", "depart": "08:30", "arrive": "11:00",
                 "duration": "2小时30分", "price": 420},
            ],
            ("北京", "大连"): [
                {"airline": "国航", "flight_no": "CA1607", "depart": "08:00", "arrive": "09:30",
                 "duration": "1小时30分", "price": 350},
                {"airline": "南航", "flight_no": "CZ6497", "depart": "14:00", "arrive": "15:30",
                 "duration": "1小时30分", "price": 300},
            ],
        }

        key = (from_city, to_city)
        results = flight_data.get(key, [])

        if not results:
            results = [{"airline": "待查询", "flight_no": "--", "depart": "--",
                        "arrive": "--", "duration": "待确认", "price": 0}]

        return results



# 方案导出为Word文档

class DocExporter:
    """将旅游方案导出为Word文档"""

    @staticmethod
    def create_travel_plan(plan_data: dict, output_path: str):
        """
        生成旅游方案Word文档
        :param plan_data: 方案数据字典
        :param output_path: 输出文件路径
        """
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # ---- 标题 ----
        title = doc.add_heading(plan_data["title"], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ---- 行程概览 ----
        doc.add_heading("行程概览", level=1)
        overview = plan_data.get("overview", {})
        for key, value in overview.items():
            p = doc.add_paragraph()
            p.add_run(f"{key}：").bold = True
            p.add_run(str(value))

        # ---- 天气与穿衣建议 ----
        if "weather" in plan_data:
            doc.add_heading("天气与穿衣建议", level=1)
            weather = plan_data["weather"]
            for city_name, weather_info in weather.items():
                doc.add_heading(city_name, level=2)
                doc.add_paragraph(
                    f"温度：{weather_info['temp_min']}°C - {weather_info['temp_max']}°C"
                )
                doc.add_paragraph(f"天气：{weather_info['desc']}")
                doc.add_paragraph(f"穿衣建议：\n{weather_info.get('clothing_advice', '')}")

        # ---- 交通方案 ----
        if "transport" in plan_data:
            doc.add_heading("交通方案", level=1)
            for segment_name, segment_info in plan_data["transport"].items():
                doc.add_heading(segment_name, level=2)
                if isinstance(segment_info, list) and len(segment_info) > 0:
                    if isinstance(segment_info[0], dict):
                        table = doc.add_table(
                            rows=1,
                            cols=len(segment_info[0]),
                            style="Table Grid"
                        )
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        # 表头
                        hdr_cells = table.rows[0].cells
                        for i, key in enumerate(segment_info[0].keys()):
                            hdr_cells[i].text = str(key)
                        # 数据行
                        for item in segment_info:
                            row_cells = table.add_row().cells
                            for i, key in enumerate(item.keys()):
                                row_cells[i].text = str(item[key])

        # ---- 酒店推荐 ----
        if "hotels" in plan_data:
            doc.add_heading("酒店推荐", level=1)
            for city_name, hotel_list in plan_data["hotels"].items():
                doc.add_heading(city_name, level=2)
                if isinstance(hotel_list, list) and len(hotel_list) > 0:
                    table = doc.add_table(
                        rows=1,
                        cols=len(hotel_list[0]),
                        style="Table Grid"
                    )
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hdr_cells = table.rows[0].cells
                    for i, key in enumerate(hotel_list[0].keys()):
                        hdr_cells[i].text = str(key)
                    for hotel in hotel_list:
                        row_cells = table.add_row().cells
                        for i, key in enumerate(hotel.keys()):
                            row_cells[i].text = str(hotel[key])

        # ---- 行程安排 ----
        if "itinerary" in plan_data:
            doc.add_heading("行程安排", level=1)
            for day_info in plan_data["itinerary"]:
                doc.add_heading(day_info.get("day", ""), level=2)
                for activity in day_info.get("activities", []):
                    p = doc.add_paragraph()
                    p.add_run(f"{activity.get('time', '')} ").bold = True
                    p.add_run(activity.get("name", ""))
                    if activity.get("score"):
                        p.add_run(f"（{activity['score']}分")
                    if activity.get("price"):
                        p.add_run(f"，{activity['price']}")
                    if activity.get("score") or activity.get("price"):
                        p.add_run("）")

        # ---- 费用预算 ----
        if "budget" in plan_data:
            doc.add_heading("费用预算", level=1)
            budget = plan_data["budget"]
            if isinstance(budget, list) and len(budget) > 0:
                table = doc.add_table(rows=1, cols=len(budget[0]), style="Table Grid")
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr_cells = table.rows[0].cells
                for i, key in enumerate(budget[0].keys()):
                    hdr_cells[i].text = str(key)
                for row in budget:
                    row_cells = table.add_row().cells
                    for i, key in enumerate(row.keys()):
                        row_cells[i].text = str(row[key])

        # ---- 温馨提示 ----
        if "tips" in plan_data:
            doc.add_heading("温馨提示", level=1)
            for tip in plan_data["tips"]:
                doc.add_paragraph(tip, style="List Bullet")

        # ---- 信息来源 ----
        if "sources" in plan_data:
            doc.add_heading("信息来源", level=1)
            doc.add_paragraph(f"数据查询时间：{plan_data['sources'].get('query_time', '')}")
            for source in plan_data["sources"].get("list", []):
                doc.add_paragraph(
                    f"{source.get('name', '')} - {source.get('content', '')} "
                    f"({source.get('timeliness', '')})",
                    style="List Bullet"
                )
            doc.add_paragraph(
                "免责声明：本攻略信息来源于公开网络渠道，仅供参考。"
                "价格信息波动较大，请以实际预订时为准。",
                style="List Bullet"
            )

        # 保存
        doc.save(output_path)
        print(f"方案已导出至：{output_path}")



# 一键生成主流程


def generate_travel_plan(
    from_city: str,
    destinations: list,
    days: int,
    budget: float = None,
    specific_needs: str = None,
    output_dir: str = None,
):
    """
    一键生成旅游方案

    :param from_city: 出发城市，如"天津"
    :param destinations: 目的地列表，如["北京", "大连"]
    :param days: 出行天数
    :param budget: 预算（元）
    :param specific_needs: 特定需求，如"环球影城、看海"
    :param output_dir: 输出目录
    :return: 方案数据字典
    """

    if output_dir is None:
        output_dir = os.path.join(os.path.expanduser("~"), "Desktop", "旅游方案")

    os.makedirs(output_dir, exist_ok=True)

    # 1. 查询天气
    print("正在查询天气...")
    weather_service = WeatherService()
    weather_data = {}
    for city in [from_city] + destinations:
        weather = weather_service.get_weather(city)
        weather["clothing_advice"] = weather_service.get_clothing_advice(weather)
        weather_data[city] = weather
        print(f"  {city}：{weather['temp_min']}°C-{weather['temp_max']}°C，{weather['desc']}")

    # 2. 查询交通
    print("正在查询交通方案...")
    transport_service = TransportService()
    transport_data = {}

    # 出发到第一个目的地
    if len(destinations) >= 1:
        trains = transport_service.search_train(from_city, destinations[0])
        flights = transport_service.search_flight(from_city, destinations[0])
        transport_data[f"{from_city} → {destinations[0]}"] = {
            "train": trains,
            "flight": flights,
        }
        print(f"  {from_city}→{destinations[0]}：高铁{len(trains)}条，飞机{len(flights)}条")

    # 目的地之间的交通
    for i in range(len(destinations) - 1):
        trains = transport_service.search_train(destinations[i], destinations[i + 1])
        flights = transport_service.search_flight(destinations[i], destinations[i + 1])
        transport_data[f"{destinations[i]} → {destinations[i + 1]}"] = {
            "train": trains,
            "flight": flights,
        }
        print(f"  {destinations[i]}→{destinations[i + 1]}：高铁{len(trains)}条，飞机{len(flights)}条")

    # 返回交通
    trains_back = transport_service.search_train(destinations[-1], from_city)
    flights_back = transport_service.search_flight(destinations[-1], from_city)
    transport_data[f"{destinations[-1]} → {from_city}（返程）"] = {
        "train": trains_back,
        "flight": flights_back,
    }

    # 3. 组装方案数据
    dest_str = "→".join(destinations)
    plan_data = {
        "title": f"{from_city}→{dest_str} {days}日游规划方案",
        "overview": {
            "出发地": from_city,
            "行程路线": f"{from_city} → {' → '.join(destinations)} → 返回{from_city}",
            "总天数": f"{days}天{days - 1}晚",
            "特定需求": specific_needs or "无",
            "预算": f"¥{budget}" if budget else "未指定",
        },
        "weather": weather_data,
        "transport": transport_data,
        "hotels": {
            city: [
                {"酒店名称": "请手动查询携程/美团填写",
                 "位置": "待确认",
                 "参考价格/晚": "待查询*",
                 "评分": "待查询",
                 "推荐理由": "建议选择交通便利、靠近景点的酒店",
                 "预订渠道": "携程/美团/飞猪"}
            ]
            for city in destinations
        },
        "itinerary": [
            {
                "day": f"第{i + 1}天",
                "activities": [
                    {"time": "上午", "name": "请根据实际需求填写景点",
                     "score": "", "price": ""},
                    {"time": "下午", "name": "请根据实际需求填写景点",
                     "score": "", "price": ""},
                    {"time": "晚上", "name": "自由活动/美食探索",
                     "score": "", "price": ""},
                ]
            }
            for i in range(days)
        ],
        "budget": [
            {"项目": "交通", "预估费用": "待计算*"},
            {"项目": "住宿", "预估费用": "待计算*"},
            {"项目": "门票", "预估费用": "待计算*"},
            {"项目": "餐饮", "预估费用": "待计算*"},
            {"项目": "总计", "预估费用": "待计算*"},
        ],
        "tips": [
            "以上价格为参考价，请以实际查询为准",
            "建议提前预订交通和住宿",
            "出行前请再次核实景点开放时间",
        ],
        "sources": {
            "query_time": datetime.now().strftime("%Y年%m月%d日"),
            "list": [
                {"name": "wttr.in", "content": "天气预报", "timeliness": "实时"},
                {"name": "12306/携程", "content": "交通时刻表及价格", "timeliness": "实时*"},
            ]
        },
    }

    # 4. 导出Word文档
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"旅游方案_{from_city}→{dest_str}_{timestamp}.docx"
    output_path = os.path.join(output_dir, filename)

    DocExporter.create_travel_plan(plan_data, output_path)

    print(f"\n方案生成完成！")
    print(f"文件路径：{output_path}")

    return plan_data


# 使用示例

if __name__ == "__main__":
    # 示例1：天津→北京→大连 3天
    generate_travel_plan(
        from_city="天津",
        destinations=["北京", "大连"],
        days=3,
        budget=6000,
        specific_needs="北京环球影城、大连看海",
    )

    # 示例2：成都→西安 3天
    # generate_travel_plan(
    #     from_city="成都",
    #     destinations=["西安"],
    #     days=3,
    #     budget=3000,
    #     specific_needs="兵马俑、回民街美食",
    # )

    # 示例3：重庆→北京→天津 5天
    # generate_travel_plan(
    #     from_city="重庆",
    #     destinations=["北京", "天津"],
    #     days=5,
    #     budget=5000,
    #     specific_needs="环球影城、煎饼果子",
    # )
